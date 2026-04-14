#!/usr/bin/env python3
"""
generate_docx.py — Générateur de document Word pour agents Microsoft 365 Copilot
Charte graphique
Usage: python generate_docx.py <config.json> <output.docx>
"""

import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installation de python-docx...")
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "--break-system-packages", "-q"])
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


# ── Palette ────────────────────────────────────────────────────────────────────
ORANGE      = RGBColor(0xFF, 0x51, 0x19)   # accent principal
DARK        = RGBColor(0x1A, 0x1A, 0x1A)   # texte courant
GRAY        = RGBColor(0x88, 0x88, 0x88)   # labels, métadonnées
LIGHT_GRAY  = RGBColor(0xF5, 0xF5, 0xF5)   # fond subtle des blocs de contenu
RULE_COLOR  = "DDDDDD"                      # couleur des séparateurs (hex str)

# ── Typographie ─────────────────────────────────────────────────────────────────
FONT_TITLE  = "Outfit"      # titres, labels, tout sauf le corps long
FONT_BODY   = "Petrona"     # corps de texte long (instructions, contenu champs)
SIZE_HERO   = 28            # pt — titre agent (page de garde)
SIZE_H1     = 16            # pt — titres de section
SIZE_H2     = 12            # pt — sous-titres / noms de champ
SIZE_BODY   = 11            # pt — contenu des champs
SIZE_META   = 9             # pt — labels, compteurs, hints
SIZE_SMALL  = 8             # pt — pied de page

# ── Noms affichables des fonctionnalités ────────────────────────────────────────
CAPABILITIES_LABELS = {
    "WebSearch":             "🔍 Recherche web",
    "OneDriveAndSharePoint": "📁 OneDrive & SharePoint",
    "Email":                 "📧 Email",
    "TeamsMessages":         "💬 Messages Teams",
    "People":                "👥 Personnes",
    "Meetings":              "📅 Réunions",
    "GraphicArt":            "🎨 Création d'images",
    "CodeInterpreter":       "💻 Interpréteur de code",
    "Dataverse":             "🗄️ Dataverse",
}


# ── Helpers XML ─────────────────────────────────────────────────────────────────

def set_run_font(run, font_name, size_pt, bold=False, color=None, italic=False):
    """Configure police, taille, gras, couleur sur un run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # Force la police sur tous les scripts
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),   font_name)
    rFonts.set(qn('w:hAnsi'),   font_name)
    rFonts.set(qn('w:eastAsia'),font_name)
    rFonts.set(qn('w:cs'),      font_name)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def add_rule(doc, before_pt=6, after_pt=12, color=RULE_COLOR):
    """Ligne horizontale fine."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before_pt)
    p.paragraph_format.space_after  = Pt(after_pt)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def set_para_shading(para, fill_hex):
    """Fond de paragraphe (hex sans #)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    existing = pPr.find(qn('w:shd'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(shd)


def add_left_border(para, color_hex="FF5119", size="12"):
    """Barre verticale orange à gauche d'un paragraphe."""
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    size)
    left.set(qn('w:space'), '12')
    left.set(qn('w:color'), color_hex)
    existing = pBdr.find(qn('w:left'))
    if existing is not None:
        pBdr.remove(existing)
    pBdr.append(left)


# ── Composants de contenu ────────────────────────────────────────────────────────

def add_section_title(doc, text, emoji=""):
    """Titre de section — Outfit, 16pt, orange, lettres capitales, séparateur dessous."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(4)
    label = f"{emoji}  {text.upper()}" if emoji else text.upper()
    run = p.add_run(label)
    set_run_font(run, FONT_TITLE, SIZE_H1, bold=True, color=ORANGE)

    # Filet orange sous le titre
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'FF5119')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_field_block(doc, field_name, content, char_limit=None, hint=None):
    """
    Bloc champ : label gris + contenu dans un encadré fond gris clair
    avec barre orange à gauche.
    """
    # ── Label ──
    label_p = doc.add_paragraph()
    label_p.paragraph_format.space_before = Pt(12)
    label_p.paragraph_format.space_after  = Pt(3)

    label_run = label_p.add_run(field_name.upper())
    set_run_font(label_run, FONT_TITLE, SIZE_H2, bold=True, color=GRAY)

    if char_limit and content:
        count = len(content)
        sep = label_p.add_run(f"   {count} / {char_limit} car.")
        set_run_font(sep, FONT_TITLE, SIZE_META, color=GRAY)

    # ── Contenu ──
    lines = (content or "(non renseigné)").split('\n')
    for i, line in enumerate(lines):
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(2) if i > 0 else Pt(0)
        cp.paragraph_format.space_after  = Pt(2)
        cp.paragraph_format.left_indent  = Cm(0.4)
        cp.paragraph_format.right_indent = Cm(0.2)

        run = cp.add_run(line if line.strip() else " ")
        is_empty = not content
        set_run_font(
            run,
            FONT_BODY,
            SIZE_BODY,
            color=GRAY if is_empty else DARK,
            italic=is_empty
        )
        set_para_shading(cp, "F5F5F5")
        add_left_border(cp)

    # ── Hint ──
    if hint:
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(3)
        hp.paragraph_format.space_after  = Pt(2)
        hp.paragraph_format.left_indent  = Cm(0.4)
        hr = hp.add_run(f"↗  {hint}")
        set_run_font(hr, FONT_TITLE, SIZE_META, color=GRAY, italic=True)


def add_suggestion_block(doc, index, titre, texte):
    """Bloc suggestion : numéro orange + titre + texte en italique."""
    # Numéro + titre
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)

    num = p.add_run(f"{index}. ")
    set_run_font(num, FONT_TITLE, SIZE_BODY, bold=True, color=ORANGE)

    title_run = p.add_run(titre)
    set_run_font(title_run, FONT_TITLE, SIZE_BODY, bold=True, color=DARK)

    # Texte
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after  = Pt(4)
    tp.paragraph_format.left_indent  = Cm(0.5)
    tr = tp.add_run(f'"{texte}"')
    set_run_font(tr, FONT_BODY, SIZE_BODY, color=GRAY, italic=True)


# ── Générateur principal ────────────────────────────────────────────────────────

def generate_document(config: dict, output_path: str):
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    agent_name = config.get("nom", "Mon Agent Copilot")
    date_str   = datetime.now().strftime("%d/%m/%Y")

    # ══════════════════════════════════════════════════
    # PAGE DE GARDE
    # ══════════════════════════════════════════════════

    # Surtitre
    sup = doc.add_paragraph()
    sup.paragraph_format.space_before = Pt(0)
    sup.paragraph_format.space_after  = Pt(6)
    sup.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sup_run = sup.add_run("MICROSOFT 365 COPILOT — AGENT BUILDER")
    set_run_font(sup_run, FONT_TITLE, SIZE_META, bold=True, color=GRAY)

    # Nom de l'agent — grand, orange
    hero = doc.add_paragraph()
    hero.paragraph_format.space_before = Pt(4)
    hero.paragraph_format.space_after  = Pt(8)
    hero.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hero_run = hero.add_run(agent_name)
    set_run_font(hero_run, FONT_TITLE, SIZE_HERO, bold=True, color=ORANGE)

    # Sous-titre
    sub = doc.add_paragraph()
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after  = Pt(6)
    sub_run = sub.add_run("Configuration complète — prêt à copier dans l'Agent Builder")
    set_run_font(sub_run, FONT_BODY, SIZE_BODY, color=GRAY)

    # Métadonnées
    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after  = Pt(16)
    meta_run = meta.add_run(
        f"Généré le {date_str}   ·   "
        "https://m365.cloud.microsoft/chat/agent/new"
    )
    set_run_font(meta_run, FONT_TITLE, SIZE_META, color=GRAY)

    add_rule(doc, before_pt=4, after_pt=20, color="FF5119")

    # ══════════════════════════════════════════════════
    # 1. IDENTITÉ
    # ══════════════════════════════════════════════════
    add_section_title(doc, "Identité de l'agent", "🤖")

    add_field_block(doc, "Nom", config.get("nom", ""),
                    char_limit=100,
                    hint="Copiez ce nom dans le champ « Nom » de l'Agent Builder.")

    add_field_block(doc, "Description", config.get("description", ""),
                    char_limit=1000,
                    hint="Cette description est visible par les utilisateurs qui parcourent la liste des agents.")

    # ══════════════════════════════════════════════════
    # 2. INSTRUCTIONS
    # ══════════════════════════════════════════════════
    add_section_title(doc, "Instructions", "📋")

    add_field_block(doc, "Instructions", config.get("instructions", ""),
                    char_limit=8000,
                    hint=(
                        "Champ le plus important. Copiez l'intégralité dans le champ « Instructions »."
                        " Affinez après les premiers tests."
                    ))

    # ══════════════════════════════════════════════════
    # 3. FONCTIONNALITÉS
    # ══════════════════════════════════════════════════
    add_section_title(doc, "Fonctionnalités à activer", "⚙️")

    caps = config.get("fonctionnalites", [])
    caps_text = "\n".join([f"✅  {CAPABILITIES_LABELS.get(c, c)}" for c in caps]) \
        if caps else "(aucune fonctionnalité spécifiée)"

    add_field_block(doc, "Fonctionnalités", caps_text,
                    hint="Activez ces fonctionnalités dans l'onglet « Fonctionnalités » de l'Agent Builder.")

    # ══════════════════════════════════════════════════
    # 4. SOURCES DE CONNAISSANCES
    # ══════════════════════════════════════════════════
    sources = config.get("sources_connaissances", [])
    if sources:
        add_section_title(doc, "Sources de connaissances", "📚")
        sources_text = "\n".join([f"•  {s}" for s in sources])
        add_field_block(doc, "Sources (URLs SharePoint / Sites web)", sources_text,
                        hint=(
                            "Ajoutez ces sources dans l'onglet « Connaissances »."
                            " Assurez-vous que les utilisateurs ont accès à ces documents."
                        ))

    # ══════════════════════════════════════════════════
    # 5. SUGGESTIONS DE DÉMARRAGE
    # ══════════════════════════════════════════════════
    suggestions = config.get("suggestions_demarrage", [])
    if suggestions:
        add_section_title(doc, "Suggestions de démarrage", "💬")

        # Compteur
        count_p = doc.add_paragraph()
        count_p.paragraph_format.space_before = Pt(2)
        count_p.paragraph_format.space_after  = Pt(8)
        count_run = count_p.add_run(f"{len(suggestions)} suggestions configurées   ·   maximum 12")
        set_run_font(count_run, FONT_TITLE, SIZE_META, color=GRAY)

        for i, sug in enumerate(suggestions, 1):
            add_suggestion_block(doc, i, sug.get("titre", f"Suggestion {i}"), sug.get("texte", ""))

    # ══════════════════════════════════════════════════
    # 6. DISCLAIMER
    # ══════════════════════════════════════════════════
    disclaimer = config.get("disclaimer", "")
    if disclaimer:
        add_section_title(doc, "Disclaimer", "⚠️")
        add_field_block(doc, "Message affiché au démarrage", disclaimer,
                        char_limit=500,
                        hint="Ce texte s'affiche à l'utilisateur au début de chaque conversation.")

    # ══════════════════════════════════════════════════
    # CHECK-LIST
    # ══════════════════════════════════════════════════
    add_rule(doc, before_pt=24, after_pt=16, color="FF5119")

    cl_title = doc.add_paragraph()
    cl_title.paragraph_format.space_before = Pt(0)
    cl_title.paragraph_format.space_after  = Pt(10)
    cl_run = cl_title.add_run("CHECK-LIST DE MISE EN LIGNE")
    set_run_font(cl_run, FONT_TITLE, SIZE_H1, bold=True, color=ORANGE)

    checklist_items = [
        ("Ouvrir l'Agent Builder",       "https://m365.cloud.microsoft/chat/agent/new"),
        ("Renseigner Nom et Description", "Champs 1 et 2"),
        ("Copier les Instructions",       "Champ le plus important — coller en intégralité"),
        ("Activer les Fonctionnalités",   "Onglet « Fonctionnalités »"),
        ("Ajouter les Sources",           "SharePoint / OneDrive — onglet « Connaissances »"),
        ("Saisir les Suggestions",        "Jusqu'à 12 suggestions de démarrage"),
        ("Tester en mode Preview",        "Posez les suggestions, affinez si besoin"),
        ("Publier et partager",           "Diffusez aux utilisateurs cibles"),
    ]

    for label, detail in checklist_items:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after  = Pt(4)
        cp.paragraph_format.left_indent  = Cm(0.2)

        checkbox = cp.add_run("☐  ")
        set_run_font(checkbox, FONT_TITLE, SIZE_BODY, bold=True, color=ORANGE)

        label_run = cp.add_run(label)
        set_run_font(label_run, FONT_TITLE, SIZE_BODY, bold=True, color=DARK)

        detail_run = cp.add_run(f"  —  {detail}")
        set_run_font(detail_run, FONT_BODY, SIZE_BODY, color=GRAY)

    # Pied de page
    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(24)
    foot.paragraph_format.space_after  = Pt(0)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_run = foot.add_run(
        f"Document généré le {date_str}   ·   Microsoft 365 Copilot Agent Builder   ·   "
        "https://m365.cloud.microsoft/chat/agent/new"
    )
    set_run_font(foot_run, FONT_TITLE, SIZE_SMALL, color=GRAY)

    # Sauvegarde
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"✅ Document généré : {output_path}")
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python generate_docx.py <config.json> <output.docx>")
        sys.exit(1)

    config_path  = sys.argv[1]
    output_path  = sys.argv[2]

    with open(config_path, "r", encoding="utf-8") as f:
        config = json.load(f)

    generate_document(config, output_path)
